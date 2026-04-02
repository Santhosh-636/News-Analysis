#!/usr/bin/env python3
"""
Convert PROJECT_REPORT.md to PROJECT_REPORT.docx
Handles markdown formatting and creates a professional Word document
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders in table"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_elem = OxmlElement(f'w:{edge}')
            edge_elem.set(qn('w:val'), 'single')
            edge_elem.set(qn('w:sz'), '12')
            edge_elem.set(qn('w:space'), '0')
            edge_elem.set(qn('w:color'), '000000')
            tcBorders.append(edge_elem)
    tcPr.append(tcBorders)

def add_horizontal_line(doc):
    """Add a horizontal line to document"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX with proper formatting"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    code_block = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            code_block = not code_block
            if code_block:
                # Start code block
                i += 1
                code_content = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_content.append(lines[i])
                    i += 1
                
                # Add code block
                if code_content:
                    code_para = doc.add_paragraph()
                    code_para.style = 'Normal'
                    code_text = '\n'.join(code_content)
                    code_run = code_para.add_run(code_text)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Set background color for code
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'E8E8E8')
                    code_para._element.get_or_add_pPr().append(shd)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        # Handle tables
        elif line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            
            # Collect table rows
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            # Parse and create table
            if len(table_lines) >= 3:  # At least header, separator, one row
                # Parse headers
                header_line = table_lines[0]
                headers = [col.strip() for col in header_line.split('|')[1:-1]]
                
                # Create table
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header
                    # Bold header
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add data rows (skip separator line)
                for line_idx in range(2, len(table_lines)):
                    if '|' in table_lines[line_idx]:
                        row_data = [col.strip() for col in table_lines[line_idx].split('|')[1:-1]]
                        if row_data and len(row_data) == len(headers):
                            row_cells = table.add_row().cells
                            for j, cell_text in enumerate(row_data):
                                row_cells[j].text = cell_text
            
            continue
        
        # Handle lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            indent_level = len(line) - len(line.lstrip())
            text = line.strip()[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        elif re.match(r'^\d+\.\s', line.strip()):
            indent_level = len(line) - len(line.lstrip())
            match = re.match(r'^(\d+)\.\s(.*)', line.strip())
            if match:
                text = match.group(2)
                doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Handle horizontal rules
        elif line.strip() in ['---', '***', '___']:
            add_horizontal_line(doc)
            i += 1
            continue
        
        # Handle empty lines
        elif not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        # Handle regular paragraphs
        else:
            # Collect paragraph lines (until empty line or special element)
            para_lines = []
            while i < len(lines) and lines[i].strip() and \
                  not lines[i].strip().startswith(('#', '|', '-', '*', '```')):
                para_lines.append(lines[i].strip())
                i += 1
            
            if para_lines:
                para_text = ' '.join(para_lines)
                # Handle bold, italic, code formatting
                para_text = para_text.replace('**', '*')  # Normalize bold
                
                p = doc.add_paragraph()
                
                # Parse inline formatting
                parts = re.split(r'(\*.*?\*|`.*?`|\[.*?\]\(.*?\))', para_text)
                for part in parts:
                    if not part:
                        continue
                    elif part.startswith('*') and part.endswith('*'):
                        # Bold
                        run = p.add_run(part[1:-1])
                        run.bold = True
                    elif part.startswith('`') and part.endswith('`'):
                        # Code
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    elif '[' in part and '(' in part:
                        # Link
                        match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                        if match:
                            run = p.add_run(match.group(1))
                            run.underline = True
                            run.font.color.rgb = RGBColor(0, 0, 255)
                    else:
                        p.add_run(part)
            
            i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✓ Document converted successfully!")
    print(f"✓ Saved as: {docx_file}")

if __name__ == '__main__':
    import os
    
    base_path = r'c:\Users\wwwsa\OneDrive\Desktop\News_Sentiment_analysis'
    md_file = os.path.join(base_path, 'PROJECT_REPORT.md')
    docx_file = os.path.join(base_path, 'PROJECT_REPORT.docx')
    
    print(f"Converting {md_file}...")
    print("This may take a moment due to the large file size...")
    
    if os.path.exists(md_file):
        markdown_to_docx(md_file, docx_file)
        print(f"\n✓ File size: {os.path.getsize(docx_file) / 1024:.1f} KB")
    else:
        print(f"Error: {md_file} not found!")
